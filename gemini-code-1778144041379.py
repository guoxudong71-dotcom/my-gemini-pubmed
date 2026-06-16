import streamlit as st
import requests
import google.generativeai as genai
import xml.etree.ElementTree as ET
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import io

# 1. 页面基本配置
st.set_page_config(page_title="BioGemini Pro - 智能文献调研站", page_icon="🧬", layout="wide")

# 2. 自定义 CSS 样式
st.markdown("""
    <style>
    .stApp { background-color: transparent; }
    .stButton>button { 
        width: 100%; border-radius: 8px; border: 1px solid #4A90E2; 
        color: #4A90E2; background-color: transparent; transition: 0.3s;
    }
    .stButton>button:hover { background-color: #4A90E2; color: white; }
    
    .paper-card { 
        padding: 24px; border-radius: 12px; background-color: rgba(128, 128, 128, 0.05); 
        border: 1px solid rgba(128, 128, 128, 0.2); margin-bottom: 20px; 
        transition: transform 0.2s ease-in-out; position: relative;
    }
    .paper-card:hover { transform: translateY(-5px); box-shadow: 0 8px 16px rgba(0,0,0,0.1); border-color: #4A90E2; }
    
    .new-tag {
        position: absolute; top: 10px; right: 10px;
        background-color: #FFD700; color: #333; padding: 2px 8px;
        border-radius: 4px; font-size: 0.75rem; font-weight: bold;
    }
    .dimension-badge {
        background-color: #2ECC71; color: white; padding: 2px 8px;
        border-radius: 4px; font-size: 0.8rem; font-weight: bold; margin-right: 10px;
    }

    .paper-title { font-size: 1.3rem; font-weight: 600; margin-bottom: 10px; line-height: 1.4; }
    .year-tag { background-color: #4A90E2; color: white; padding: 2px 10px; border-radius: 4px; font-size: 0.85rem; font-weight: bold; margin-right: 10px; }
    .ai-box { background-color: rgba(74, 144, 226, 0.1); padding: 25px; border-left: 6px solid #4A90E2; border-radius: 8px; margin: 15px 0; width: 100%; line-height: 1.7; }
    .report-box { background-color: rgba(46, 204, 113, 0.08); padding: 25px; border-left: 6px solid #2ECC71; border-radius: 8px; margin: 20px 0; line-height: 1.8; }
    </style>
    """, unsafe_allow_html=True)

# 3. 配置 Gemini
api_key = st.secrets.get("GEMINI_API_KEY")
model = None
if api_key:
    try:
        genai.configure(api_key=api_key)
        available_models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
        selected_model = next((m for m in available_models if 'gemini-1.5-flash' in m), available_models[0])
        model = genai.GenerativeModel(selected_model)
        st.sidebar.success(f"✅ 已连接: {selected_model}")
    except Exception as e:
        st.sidebar.error(f"❌ 初始化失败: {e}")

# 4. 核心数据抓取函数
def search_pubmed_advanced(query, years=5, max_results=10, sort="relevance", days_limit=None):
    search_url = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esearch.fcgi"
    if days_limit:
        advanced_query = f"{query} AND (\"last {days_limit} days\"[Filter])"
    else:
        current_year = datetime.now().year
        min_year = current_year - years
        advanced_query = f"({query}) AND ({min_year}:{current_year}[DP])"
    
    params = {"db": "pubmed", "term": advanced_query, "retmax": max_results, "retmode": "json", "sort": sort}
    try:
        r = requests.get(search_url, params=params, timeout=10)
        return r.json().get("esearchresult", {}).get("idlist", [])
    except: return []

def get_details(pmid):
    fetch_url = f"https://eutils.ncbi.nlm.nih.gov/entrez/eutils/efetch.fcgi?db=pubmed&id={pmid}&retmode=xml"
    try:
        r = requests.get(fetch_url, timeout=10)
        root = ET.fromstring(r.content)
        title = root.find(".//ArticleTitle").text if root.find(".//ArticleTitle") is not None else "Untitled"
        abstracts = root.findall(".//AbstractText")
        abstract_text = " ".join([n.text for n in abstracts if n.text])
        pub_year = root.find(".//PubDate/Year")
        year = pub_year.text if pub_year is not None else "Recent"
        return title, abstract_text, year
    except: return None, None, None

# 导出格式化 Word 报告的辅助函数
def create_topic_report_docx(topic, report_text, text_sources):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.font._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    doc.add_heading(f'【专题学术简报】{topic} 领域最新研究进展', 0)
    doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d')}  |  数据源: PubMed 实时数据库")
    
    doc.add_heading('一、核心进展综述汇总', level=1)
    for line in report_text.split('\n'):
        if line.strip():
            para = doc.add_paragraph(line)
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                
    doc.add_heading('二、本次追踪参考的原始文献', level=1)
    for idx, src in enumerate(text_sources, 1):
        p = doc.add_paragraph()
        p.add_run(f"[{idx}] ").bold = True
        p.add_run(f"标题: {src['title']}\n    PMID: {src['pmid']} ({src['year']})")
        
    doc.add_paragraph('\n--- Generated by BioGemini Pro ---')
    bio = io.BytesIO(); doc.save(bio); return bio.getvalue()

# 5. 界面布局与控制台
with st.sidebar:
    st.markdown("### 🧬 控制中心")
    mode = st.radio("工作模式选择", ["🎯 专案主题追踪", "🔍 常规文献检索"])
    
    st.markdown("---")
    is_tracking = st.toggle("🛰️ 仅看最新发布(追新)")
    track_days = st.select_slider("追新窗口", options=[1, 3, 7, 30, 90], value=30, format_func=lambda x: f"最近 {x} 天") if is_tracking else None
    
    if mode == "🔍 常规文献检索":
        search_years = st.slider("常规时间跨度 (年)", 1, 20, 5)
        search_sort = st.selectbox("排序规则", ["relevance", "pub_date"], format_func=lambda x: "相关性优先" if x=="relevance" else "最新日期优先")
    max_num = st.number_input("各维度最大解析篇数", 1, 20, 3 if mode == "🎯 专案主题追踪" else 10)

st.title("🔬 BioGemini Pro")
st.markdown("##### 结合 PubMed 实时检索与 Gemini 1.5 深度分析的学术助手")

# ==================== 模式一：🎯 专案主题追踪 ====================
if mode == "🎯 专案主题追踪":
    st.markdown("### 📋 领域多维度自动追踪与汇总")
    col1, col2 = st.columns([3, 1])
    with col1:
        topic_input = st.text_input("请输入你想重点监控的核心领域词：", value="RSV", placeholder="例如：RSV、Alzheimer、mRNA Vaccine...")
    with col2:
        st.markdown("<div style='margin-top:28px;'></div>", unsafe_allow_html=True)
        generate_report_btn = st.button("🚀 一键生成多维度汇总研报")

    # 定义要追踪的核心维度及对应的 PubMed 检索补充词
    dimensions = {
        "疾病负担 (Burden of Disease)": "disease burden OR economic burden OR mortality",
        "发病率与流行病学 (Incidence & Epidemiology)": "incidence OR prevalence OR epidemiology",
        "预防与疫苗新进展 (Prevention & Vaccines)": "prevention OR vaccine OR prophylaxis OR monoclonal antibody"
    }

    if generate_report_btn and topic_input:
        all_fetched_data = []
        raw_papers_for_word = []
        
        # 依次抓取各个维度的文献
        for dim_name, dim_query in dimensions.items():
            combined_query = f"({topic_input}) AND ({dim_query})"
            with st.spinner(f"正在检索【{dim_name}】相关文献..."):
                ids = search_pubmed_advanced(combined_query, years=5, max_results=max_num, sort="pub_date" if is_tracking else "relevance", days_limit=track_days)
            
            dim_papers = []
            for pmid in ids:
                title, abstract, year = get_details(pmid)
                if title and abstract:
                    dim_papers.append({"title": title, "abstract": abstract, "year": year, "pmid": pmid})
                    raw_papers_for_word.append({"title": title, "year": year, "pmid": pmid})
            
            if dim_papers:
                all_fetched_data.append({"dimension": dim_name, "papers": dim_papers})

        # 如果抓到了数据，交给 Gemini 进行交叉大融合
        if all_fetched_data and model:
            st.success("📝 数据抓取完毕！AI 正在跨文献提炼、汇总核心要点...")
            
            # 组装超级 Prompt
            prompt = f"你是一位世界顶尖的生物医学情报专家。请根据我提供的 PubMed 最新文献摘要，针对【{topic_input}】领域，撰写一份高度整合、逻辑严密的学术进展简报。不要单篇罗列，要分维度进行多文献的交叉归纳总结。\n\n"
            for block in all_fetched_data:
                prompt += f"### 维度：{block['dimension']}\n"
                for i, p in enumerate(block['papers'], 1):
                    prompt += f"文献{i} Title: {p['title']} (PMID: {p['pmid']})\nAbstract: {p['abstract']}\n\n"
            
            prompt += "请严格按以下结构用中文输出：\n1. 【该领域近期动态大势综述】\n2. 【疾病负担与发病率核心数据/结论整合】\n3. 【预防、疫苗或临床干预的最新突破与技术路线】\n4. 【目前研究存在的空白或下一步方向】"
            
            try:
                with st.spinner("Gemini 正在全速撰写简报中..."):
                    response = model.generate_content(prompt)
                    report_content = response.text
                
                # 界面展示 AI 汇总报告
                st.markdown("### 📊 AI 多文献交叉智能化研报")
                st.markdown(f'<div class="report-box">{report_content}</div>', unsafe_allow_html=True)
                
                # 转换为完美字体的 Word 报告提供下载
                report_docx = create_topic_report_docx(topic_input, report_content, raw_papers_for_word)
                st.download_button(
                    label="📥 下载这份主题整合 Word 报告",
                    data=report_docx,
                    file_name=f"{topic_input}_Status_Report_{datetime.now().strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
            except Exception as e:
                st.error(f"AI 研报生成失败: {e}")
        else:
            st.warning("在此检索条件下未找到足够文献，请尝试放宽追踪窗口或更改关键词。")

# ==================== 模式二：🔍 常规文献检索 ====================
elif mode == "🔍 常规文献检索":
    user_query = st.text_input("", placeholder="输入研究关键词 (例如: RSV prevention)...", label_visibility="collapsed")

    if user_query:
        with st.spinner("🚀 正在为您扫描数据库..."):
            ids = search_pubmed_advanced(user_query, years=search_years, max_results=max_num, sort=search_sort, days_limit=track_days)
        
        if ids:
            st.success(f"已为您精选 {len(ids)} 篇文献")
            for pmid in ids:
                title, abstract, year = get_details(pmid)
                if not title: continue
                
                tag_html = '<div class="new-tag">NEW</div>' if is_tracking else ''
                card_html = f'<div class="paper-card">{tag_html}<span class="year-tag">{year}</span><span style="opacity: 0.7; font-size: 0.85rem;">PMID: {pmid}</span><div class="paper-title">{title}</div><div style="margin-top: 10px;"><a href="https://pubmed.ncbi.nlm.nih.gov/{pmid}/" target="_blank" style="text-decoration: none; color: #4A90E2;">🔗 查看原文 (PubMed)</a></div></div>'
                st.markdown(card_html, unsafe_allow_html=True)
                
                btn_col, _ = st.columns([1.5, 5])
                with btn_col:
                    analyze_btn = st.button("✨ 深度分析摘要", key=f"ai_{pmid}")
                
                if analyze_btn:
                    if model:
                        with st.spinner("AI 正在深度研读中..."):
                            prompt = f"针对以下摘要进行深度分析并用中文回答：1.【中文标题翻译】 2.【核心结论】 3.【亮点与局限】。内容：{abstract}"
                            try:
                                response = model.generate_content(prompt)
                                ai_content = response.text
                                st.markdown(f'<div class="ai-box">{ai_content}</div>', unsafe_allow_html=True)
                                
                                # 这里保留了之前完全修好的中西文字体一致的单篇导出
                                docx_file = create_topic_report_docx(title[:20], f"针对单篇文献的深度分析：\n\n{ai_content}", [{"title": title, "pmid": pmid, "year": year}])
                                st.download_button("📥 下载 Word 报告", docx_file, f"Analysis_{pmid}.docx", key=f"dl_{pmid}")
                            except Exception as e: st.error(f"分析失败: {e}")
                st.markdown("<hr style='opacity: 0.1; margin: 20px 0;'>", unsafe_allow_html=True)
        else:
            st.warning("暂未发现相关文献。")

st.markdown("---")
st.caption("© 2026 BioGemini | 自动追踪与多维度智能研报")
